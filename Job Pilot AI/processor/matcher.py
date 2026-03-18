import re
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity

SKILLS = ["python","java","javascript","typescript","react","angular","vue","node.js",
          "django","flask","fastapi","spring","sql","mysql","postgresql","mongodb",
          "redis","aws","azure","gcp","docker","kubernetes","git","linux",
          "machine learning","deep learning","tensorflow","pytorch","pandas",
          "numpy","scikit-learn","nlp","data science","spark","tableau","power bi",
          "c++","c#","golang","rust","php","ruby","scala","rest api","graphql",
          "microservices","agile","scrum","ci/cd","jenkins","terraform"]

def extract_skills(text):
    tl = text.lower()
    return [s for s in SKILLS if s in tl]

def parse_resume(file_path):
    """Parse PDF or DOCX resume and return text + skills."""
    text = ''
    ext  = file_path.lower().split('.')[-1]
    try:
        if ext == 'pdf':
            import pdfplumber
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    text += page.extract_text() or ''
        elif ext in ['docx', 'doc']:
            from docx import Document
            doc = Document(file_path)
            for para in doc.paragraphs:
                text += para.text + '\n'
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + ' '
        else:
            return {'text': '', 'skills': [], 'word_count': 0,
                    'error': f'Unsupported file type: {ext}'}

        text = re.sub(r'\s+', ' ', text.strip())
        return {
            'text':       text,
            'skills':     extract_skills(text),
            'word_count': len(text.split())
        }
    except Exception as e:
        return {'text': '', 'skills': [], 'word_count': 0, 'error': str(e)}

def score_match(resume_text, job_desc, job_skills=''):
    if not resume_text: return 0.0
    try:
        docs  = [resume_text, job_desc + ' ' + job_skills]
        tfidf = TfidfVectorizer(stop_words='english').fit_transform(docs)
        return round(float(cosine_similarity(tfidf[0:1], tfidf[1:2])[0][0]), 3)
    except Exception:
        if not job_skills: return 0.0
        skills  = [s.strip().lower() for s in job_skills.split(',') if s.strip()]
        matched = sum(1 for s in skills if s in resume_text.lower())
        return round(matched / len(skills), 3) if skills else 0.0
